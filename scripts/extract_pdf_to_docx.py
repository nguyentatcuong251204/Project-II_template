from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document

PDF_PATH = Path('main.pdf')
OUT_DOCX = Path('Do_an_template_from_main.docx')

if not PDF_PATH.exists():
    print(f"Error: {PDF_PATH} not found in working directory.")
    raise SystemExit(1)

reader = PdfReader(str(PDF_PATH))
doc = Document()

for i, page in enumerate(reader.pages):
    try:
        text = page.extract_text() or ''
    except Exception as e:
        text = ''
    if text.strip():
        # Remove problematic control characters that break python-docx/lxml
        def sanitize(s: str) -> str:
            out_chars = []
            for ch in s:
                cp = ord(ch)
                # allow common whitespace and printable Unicode ranges
                if ch in ('\\n', '\\t', '\\r'):
                    out_chars.append(ch)
                elif 32 <= cp <= 0xD7FF or 0xE000 <= cp <= 0xFFFD:
                    out_chars.append(ch)
                else:
                    out_chars.append(' ')
            return ''.join(out_chars).replace('\\r', '')

        clean = sanitize(text)
        # Simple normalization: split on double newlines or single newlines
        paras = [p.strip() for p in clean.split('\\n\\n') if p.strip()]
        if not paras:
            paras = [p.strip() for p in clean.split('\\n') if p.strip()]
        for p in paras:
            doc.add_paragraph(p)
    else:
        doc.add_paragraph('')
    if i != len(reader.pages) - 1:
        doc.add_page_break()

doc.save(str(OUT_DOCX))
print(f"Saved extracted content to {OUT_DOCX}")
